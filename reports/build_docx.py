"""Build a .docx version of Reflective_Synthesis_Paper.md so the user can
adjust page starts (insert /Ctrl+Enter/) per heading.

Why this exists: the PDF version is a single flowing render; Word lets the
user decide where each major section starts. We mirror the PDF's typography
(11pt Segoe UI body, dark-blue H1/H2 underline) so the two stay in sync.

Notes on the mermaid diagram: python-docx cannot render mermaid to SVG
without a heavy dep. We emit a labelled placeholder paragraph (Heading 4
style, light grey border) that the user can replace with a paste of
`reports/agent_graph.png` or a screenshot from the .html. The .html is
deleted by the PDF builder unless --keep-html is passed; rerun with
--keep-html if you need a fresh screenshot source.

Usage:
    python reports/build_docx.py
Outputs:
    reports/Reflective_Synthesis_Paper.docx
"""
from __future__ import annotations
import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm, Inches, Emu

_REPORT_MD = Path(__file__).resolve().parent / "Reflective_Synthesis_Paper.md"
_DOCX_OUT  = Path(__file__).resolve().parent / "Reflective_Synthesis_Paper.docx"

# Same body colour as the PDF.
_BLUE = RGBColor(0x0B, 0x3D, 0x91)
_LIGHT_GREY = RGBColor(0x55, 0x55, 0x55)
_BG_GREY = "F3F5F8"
_BORDER_GREY = "D6DDE6"


# ---------------------------------------------------------------------------
# Markdown -> a list of typed blocks. We do this in two stages: parse to
# blocks, then walk blocks and emit Word paragraphs/runs. This keeps the
# PDF builder's _md_to_html simple (string concat) and the docx builder
# style-aware (no HTML string munging).
# ---------------------------------------------------------------------------

def _parse_blocks(md: str) -> list[dict]:
    """Parse markdown into a flat list of typed blocks.

    Block types: 'h1', 'h2', 'h3', 'p', 'ul', 'ol', 'code', 'mermaid',
    'hr', 'blockquote'.

    Inline structure (bold/italic/code/sup/links) is stored as a list of
    'inline' dicts under 'inlines'; each is {text, bold, italic, code, sup}.
    """
    blocks: list[dict] = []
    lines = md.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        s = line.strip()

        # Fenced code.
        if s.startswith("```"):
            lang = s[3:].strip()
            i += 1
            buf: list[str] = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                buf.append(lines[i]); i += 1
            i += 1  # closing fence
            if lang == "mermaid":
                blocks.append({"type": "mermaid", "text": "\n".join(buf)})
            else:
                blocks.append({"type": "code", "text": "\n".join(buf)})
            continue

        # Headings.
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            blocks.append({"type": f"h{min(level, 3)}", "inlines": _parse_inlines(m.group(2))})
            i += 1; continue

        # Horizontal rule.
        if re.match(r"^-{3,}\s*$", s):
            blocks.append({"type": "hr"}); i += 1; continue

        # Blockquote.
        if s.startswith(">"):
            content = s.lstrip("> ").strip()
            blocks.append({"type": "blockquote", "inlines": _parse_inlines(content)})
            i += 1; continue

        # Unordered list.
        if re.match(r"^[-*]\s+", line):
            items: list[list[dict]] = []
            while i < len(lines) and re.match(r"^[-*]\s+", lines[i]):
                items.append(_parse_inlines(re.sub(r"^[-*]\s+", "", lines[i])))
                i += 1
            blocks.append({"type": "ul", "items": items})
            continue

        # Ordered list.
        if re.match(r"^\d+\.\s+", line):
            items = []
            while i < len(lines) and re.match(r"^\d+\.\s+", lines[i]):
                items.append(_parse_inlines(re.sub(r"^\d+\.\s+", "", lines[i])))
                i += 1
            blocks.append({"type": "ol", "items": items})
            continue

        # Blank line.
        if not s:
            i += 1; continue

        # Plain paragraph: gather until the next block boundary.
        para = [line]; i += 1
        while i < len(lines):
            nxt = lines[i]
            ns = nxt.strip()
            if (not ns
                or ns.startswith("```")
                or re.match(r"^#{1,6}\s+", nxt)
                or re.match(r"^-{3,}\s*$", ns)
                or ns.startswith(">")
                or re.match(r"^[-*]\s+", nxt)
                or re.match(r"^\d+\.\s+", nxt)):
                break
            para.append(nxt); i += 1
        blocks.append({"type": "p", "inlines": _parse_inlines(" ".join(para))})
    return blocks


def _parse_inlines(s: str) -> list[dict]:
    """Parse inline markdown into a list of {text, bold, italic, code, sup} dicts.

    Preserves raw HTML (e.g. the paper's <sup>1</sup> citation markers) as
    a 'sup' inline type so we can emit it as a real Word superscript run.
    """
    # Split on HTML <sup>...</sup> tokens first; we treat them as one
    # inline-typed 'sup' block each.
    out: list[dict] = []
    pos = 0
    sup_re = re.compile(r"<sup>(.*?)</sup>", re.DOTALL)
    for m in sup_re.finditer(s):
        if m.start() > pos:
            out.extend(_parse_md_inlines(s[pos:m.start()]))
        out.append({"text": m.group(1), "sup": True})
        pos = m.end()
    if pos < len(s):
        out.extend(_parse_md_inlines(s[pos:]))
    return out


def _parse_md_inlines(s: str) -> list[dict]:
    """Parse **bold**, *italic*, `code`, and [text](url) inside a span
    that has no HTML. Greedy left-to-right tokenisation.
    """
    # Order matters: bold (longest), then italic, then inline code, then link.
    pattern = re.compile(
        r"\*\*(.+?)\*\*"            # bold
        r"|\*(?!\s)([^*]+?)\*(?!\*)"  # italic
        r"|`([^`]+)`"                # code
        r"|\[([^\]]+)\]\(([^)]+)\)"  # link
    )
    out: list[dict] = []
    pos = 0
    for m in pattern.finditer(s):
        if m.start() > pos:
            out.append({"text": s[pos:m.start()]})
        if m.group(1) is not None:     # bold
            out.append({"text": m.group(1), "bold": True})
        elif m.group(2) is not None:   # italic
            out.append({"text": m.group(2), "italic": True})
        elif m.group(3) is not None:   # code
            out.append({"text": m.group(3), "code": True})
        elif m.group(4) is not None:   # link
            out.append({"text": m.group(4), "url": m.group(5)})
        pos = m.end()
    if pos < len(s):
        out.append({"text": s[pos:]})
    return out


# ---------------------------------------------------------------------------
# docx construction
# ---------------------------------------------------------------------------

def _set_cell_border(cell, **kwargs) -> None:
    """Apply borders to a single table cell. Used for the mermaid placeholder
    and code blocks rendered as 1x1 tables (to get the grey border reliably
    in Word — paragraph borders are fiddly across Word versions)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        if edge in kwargs:
            tag = OxmlElement(f"w:{edge}")
            tag.set(qn("w:val"), kwargs[edge].get("val", "single"))
            tag.set(qn("w:sz"),  str(kwargs[edge].get("sz", 4)))
            tag.set(qn("w:color"), kwargs[edge].get("color", _BORDER_GREY))
            tcBorders.append(tag)
    tcPr.append(tcBorders)


def _shade_cell(cell, fill: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def _add_page_number_field(paragraph) -> None:
    """Insert a Word PAGE / NUMPAGES field so the footer reads 'N / M'.

    PAGE and NUMPAGES are live fields: Word updates them on open / print.
    We insert fldChar begin / instrText / fldChar end as the canonical
    pattern for live fields, separated by a literal ' / '.
    """
    def _fld(text: str) -> None:
        run = paragraph.add_run()
        r = run._r
        fld_begin = OxmlElement("w:fldChar"); fld_begin.set(qn("w:fldCharType"), "begin")
        r.append(fld_begin)
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = f" {text} "
        r.append(instr)
        fld_sep = OxmlElement("w:fldChar"); fld_sep.set(qn("w:fldCharType"), "separate")
        r.append(fld_sep)
        # A literal placeholder so Word has something to show before the
        # field is updated. Single space keeps the layout stable.
        t = OxmlElement("w:t"); t.text = "1"; r.append(t)
        fld_end = OxmlElement("w:fldChar"); fld_end.set(qn("w:fldCharType"), "end")
        r.append(fld_end)

    _fld("PAGE")
    paragraph.add_run(" / ")
    _fld("NUMPAGES")


def _add_run(p, inline: dict, base_font: str = "Segoe UI", base_size: int = 11) -> None:
    """Add one run to paragraph `p` from an inline dict."""
    text = inline.get("text", "")
    if not text:
        return
    r = p.add_run(text)
    r.font.name = base_font
    r.font.size = Pt(base_size)
    if inline.get("bold"): r.bold = True
    if inline.get("italic"): r.italic = True
    if inline.get("sup"):
        r.font.superscript = True
        r.font.size = Pt(max(7, base_size - 4))  # 7pt looks right at body=11
        r.font.color.rgb = _BLUE
    if inline.get("code"):
        r.font.name = "Cascadia Mono"
        r.font.size = Pt(10)
    if "url" in inline:
        # python-docx's add_hyperlink is not built-in; we just style it as a
        # blue underline run so the URL is visible after a paste. A real
        # hyperlink would need raw XML; skip for a v1 builder.
        r.font.color.rgb = _BLUE
        r.font.underline = True


def _set_para_margins(p, before: int = 0, after: int = 0, line: float | None = None) -> None:
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line is not None:
        pf.line_spacing = line


def _ensure_styles(doc: Document) -> None:
    """Tune the built-in Heading 1/2/3 + Normal styles to match the PDF.

    A4 page, ~18mm margins (matches the PDF @page), dark-blue headings,
    light-grey borders under H1/H2.
    """
    # Page setup on the (single) section.
    s = doc.sections[0]
    s.page_width  = Cm(21.0)
    s.page_height = Cm(29.7)
    s.top_margin    = Cm(1.8)
    s.bottom_margin = Cm(2.2)
    s.left_margin   = Cm(1.6)
    s.right_margin  = Cm(1.6)
    s.header_distance = Cm(1.0)
    s.footer_distance = Cm(1.0)

    # Body / Normal.
    normal = doc.styles["Normal"]
    normal.font.name = "Segoe UI"
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(6)

    # Headings.
    h1 = doc.styles["Heading 1"]
    h1.font.name = "Segoe UI"; h1.font.size = Pt(19)
    h1.font.bold = True; h1.font.color.rgb = _BLUE
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after  = Pt(4)
    h1.paragraph_format.keep_with_next = True

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Segoe UI"; h2.font.size = Pt(14)
    h2.font.bold = True; h2.font.color.rgb = _BLUE
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after  = Pt(3)
    h2.paragraph_format.keep_with_next = True

    h3 = doc.styles["Heading 3"]
    h3.font.name = "Segoe UI"; h3.font.size = Pt(12)
    h3.font.bold = True; h3.font.color.rgb = _BLUE
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after  = Pt(2)
    h3.paragraph_format.keep_with_next = True


def _add_footer(doc: Document) -> None:
    footer = doc.sections[0].footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # Clear any default run so the field is the only content.
    for r in list(p.runs):
        r.text = ""
    _add_page_number_field(p)


def _add_h1_with_border(p) -> None:
    """Add a thin blue border under a Heading 1 paragraph. python-docx
    doesn't expose paragraph borders on a style directly, so we apply a
    bottom border to the specific paragraph."""
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")          # 1pt
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "0B3D91")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _emit_block(doc: Document, block: dict) -> None:
    t = block["type"]
    if t == "h1":
        p = doc.add_paragraph(style="Heading 1")
        for inl in block["inlines"]: _add_run(p, inl)
        _add_h1_with_border(p)
    elif t == "h2":
        p = doc.add_paragraph(style="Heading 2")
        for inl in block["inlines"]: _add_run(p, inl)
    elif t == "h3":
        p = doc.add_paragraph(style="Heading 3")
        for inl in block["inlines"]: _add_run(p, inl)
    elif t == "p":
        p = doc.add_paragraph()
        for inl in block["inlines"]: _add_run(p, inl)
    elif t == "ul":
        for items in block["items"]:
            p = doc.add_paragraph(style="List Bullet")
            for inl in items: _add_run(p, inl)
    elif t == "ol":
        for items in block["items"]:
            p = doc.add_paragraph(style="List Number")
            for inl in items: _add_run(p, inl)
    elif t == "blockquote":
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.right_indent = Cm(0.6)
        for inl in block["inlines"]:
            _add_run(p, inl, base_size=11)
            # Lighten the quote.
            for r in p.runs:
                r.font.color.rgb = _LIGHT_GREY
    elif t == "hr":
        # Word has a built-in horizontal rule via inserting an empty
        # paragraph with a bottom border. Simpler: use a separator character.
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        r = p.add_run("─" * 30)
        r.font.color.rgb = RGBColor(0xCF, 0xD8, 0xE3)
    elif t == "code":
        # Render code in a 1x1 shaded table so the grey box + border look
        # right in Word. Preserve line breaks and indentation.
        text = block["text"]
        table = doc.add_table(rows=1, cols=1)
        table.autofit = True
        cell = table.rows[0].cells[0]
        _shade_cell(cell, _BG_GREY)
        _set_cell_border(cell, top={"sz": 4, "color": _BORDER_GREY},
                              left={"sz": 4, "color": _BORDER_GREY},
                              bottom={"sz": 4, "color": _BORDER_GREY},
                              right={"sz": 4, "color": _BORDER_GREY})
        # Replace the auto-created empty paragraph with our code text.
        cell.paragraphs[0].text = ""
        for j, line in enumerate(text.splitlines()):
            p = cell.paragraphs[0] if j == 0 else cell.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(line if line else " ")
            r.font.name = "Cascadia Mono"
            r.font.size = Pt(10)
        # A small spacer paragraph so the next block doesn't touch the table.
        doc.add_paragraph()
    elif t == "mermaid":
        # Honest placeholder. We do not render the diagram here; the user
        # pastes the screenshot (e.g. reports/agent_graph.png) into the
        # placeholder box, or replaces this block with a Word-native shape.
        table = doc.add_table(rows=2, cols=1)
        table.autofit = True
        title_cell = table.rows[0].cells[0]
        body_cell  = table.rows[1].cells[0]
        for c in (title_cell, body_cell):
            _set_cell_border(c, top={"sz": 4, "color": _BORDER_GREY},
                                  left={"sz": 4, "color": _BORDER_GREY},
                                  bottom={"sz": 4, "color": _BORDER_GREY},
                                  right={"sz": 4, "color": _BORDER_GREY})
        _shade_cell(title_cell, _BG_GREY)
        title_cell.paragraphs[0].text = ""
        r = title_cell.paragraphs[0].add_run("[ Architecture diagram — paste screenshot here ]")
        r.bold = True; r.font.name = "Segoe UI"; r.font.size = Pt(11)
        r.font.color.rgb = _BLUE
        body_cell.paragraphs[0].text = ""
        # Show the mermaid source as monospace text inside the body cell so
        # the user has the source on hand when re-rendering for the screenshot.
        for j, line in enumerate(block["text"].splitlines()):
            p = body_cell.paragraphs[0] if j == 0 else body_cell.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(line if line else " ")
            r.font.name = "Cascadia Mono"; r.font.size = Pt(8)
            r.font.color.rgb = _LIGHT_GREY
        doc.add_paragraph()  # spacer
    else:
        # Defensive: emit as a plain paragraph.
        p = doc.add_paragraph()
        p.add_run(str(block))


def main() -> None:
    ap = argparse.ArgumentParser(description="Render the synthesis paper to a .docx.")
    args = ap.parse_args()

    md = _REPORT_MD.read_text(encoding="utf-8")
    blocks = _parse_blocks(md)

    doc = Document()
    _ensure_styles(doc)
    _add_footer(doc)

    for b in blocks:
        _emit_block(doc, b)

    doc.save(_DOCX_OUT)
    print(f"wrote docx: {_DOCX_OUT}")


if __name__ == "__main__":
    main()
